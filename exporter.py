from docx import Document
from docx.shared import Pt

def export_to_docx(chat_log, filename="Business_Plan.docx"):
    doc = Document()
    doc.add_heading("Business Plan", 0)

    section_map = {
        "Business Info": ["business name", "type of business", "services", "customers"],
        "SWOT Analysis": ["strengths", "weaknesses", "opportunities", "threats", "SWOT"],
        "Vision, Mission & Objectives": ["vision", "mission", "objectives"],
        "Strategies": ["strategy"],
        "Marketing Plan": ["marketing mix", "target market", "positioning", "marketing objectives"],
        "Executive Summary": ["executive summary"]
    }

    sections = {section: [] for section in section_map}

    for entry in chat_log:
        if entry["role"] != "assistant":
            continue
        content = entry["content"].lower()
        for section, keywords in section_map.items():
            if any(keyword in content for keyword in keywords):
                sections[section].append(entry["content"])
                break

    for section, texts in sections.items():
        if texts:
            doc.add_heading(section, level=1)
            for t in texts:
                doc.add_paragraph(t).style.font.size = Pt(11)

    doc.save(filename)
